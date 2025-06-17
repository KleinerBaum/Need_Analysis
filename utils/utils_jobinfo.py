# utils_jobinfo.py
"""Utility functions for job data extraction and session management."""

from __future__ import annotations

import base64
import re
from typing import Dict, Optional

from utils.keys import ALL_STEP_KEYS

from utils.i18n import tr

import docx
import fitz  # PyMuPDF
import streamlit as st


def extract_text_from_pdf(file) -> str:
    """Return plain text from a PDF file."""
    file.seek(0)
    data = file.read()
    doc = fitz.open(stream=data, filetype="pdf")
    return "".join(page.get_text() for page in doc)


def extract_text_from_docx(file) -> str:
    """Return text from a DOCX file including tables."""

    file.seek(0)
    doc = docx.Document(file)
    lines = [para.text for para in doc.paragraphs]

    # Include table cell text (if any). Some job ad templates store
    # important fields like the job title in tables.
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text:
                    lines.append(cell.text)

    return "\n".join(lines)


def detect_file_type(file) -> Optional[str]:
    """Detect file type based on extension."""
    name = file.name.lower()
    if name.endswith(".pdf"):
        return "pdf"
    if name.endswith(".docx"):
        return "docx"
    if name.endswith(".txt"):
        return "txt"
    return None


def extract_text(file) -> str:
    """Extract content from a supported file."""
    filetype = detect_file_type(file)
    if filetype == "pdf":
        return extract_text_from_pdf(file)
    if filetype == "docx":
        return extract_text_from_docx(file)
    if filetype == "txt":
        file.seek(0)
        return file.read().decode("utf-8")
    raise ValueError("Unsupported file type")


def basic_field_extraction(text: str) -> Dict[str, str]:
    """Return a dictionary with extracted fields from ``text``.

    This naive regex approach detects ``job_title``, ``company_name`` and some
    simple skill statements. The raw text is stored under ``parsed_data_raw``.
    Any missing keys from :data:`keys.ALL_STEP_KEYS` are included with empty
    strings so that Streamlit widgets can be pre-populated consistently.
    """

    fields: Dict[str, str] = {"parsed_data_raw": text}

    job_title = re.search(
        r"(?im)^\s*(Stellenbezeichnung|Jobtitel|Position)\s*[:\-]\s*(.+)$",
        text,
    )
    if job_title:
        fields["job_title"] = job_title.group(2).strip()

    company_name = re.search(
        r"(?im)^\s*(Unternehmen|Company|Firma)\s*[:\-]\s*(.+)$",
        text,
    )
    if company_name:
        fields["company_name"] = company_name.group(2).strip()

    city = re.search(
        r"(?im)^\s*(Stadt|Ort|City)\s*[:\-]\s*(.+)$",
        text,
    )
    if city:
        fields["city"] = city.group(2).strip()

    website = re.search(
        r"(?im)^\s*(Unternehmenswebsite|Website|Webseite|Company Website)\s*[:\-]\s*(\S+)",
        text,
    )
    if website:
        fields["company_website"] = website.group(2).strip()

    job_type_match = re.search(
        r"(?i)\b(full[-\s]?time|teilzeit|part[-\s]?time|praktikum|internship|freelance)\b",
        text,
    )
    if job_type_match:
        fields["job_type"] = job_type_match.group(1).replace("-", " ").title()

    contract_type_match = re.search(
        r"(?i)\b(unbefristet|permanent|befristet|fixed[-\s]?term|werkvertrag|contract for work)\b",
        text,
    )
    if contract_type_match:
        fields["contract_type"] = contract_type_match.group(1).replace("-", " ").title()

    job_level_match = re.search(r"(?i)\b(junior|mid|senior|lead|management)\b", text)
    if job_level_match:
        fields["job_level"] = job_level_match.group(1).title()

    # --- very small skill extraction -------------------------------------
    cleaned_text = re.sub(r"e\.g\.,?", "", text, flags=re.IGNORECASE)
    cleaned_text = re.sub(r"i\.e\.,?", "", cleaned_text, flags=re.IGNORECASE)
    skill_phrases = re.findall(
        r"(?i)(?:proficiency in|experience with|knowledge of|proficient in)\s+(.+?)(?:\.|\n)",
        cleaned_text,
    )
    skills: list[str] = []
    for phrase in skill_phrases:
        clean = re.sub(r"\(e\.g\.,?", "", phrase)
        clean = re.sub(r"[()]", ",", clean)
        clean = re.sub(r"[\.\n]", "", clean)
        clean = re.sub(r"\s+", " ", clean)
        parts = re.split(r",| and | und |&", clean)
        for part in parts:
            part = part.strip()
            if part and part not in skills:
                skills.append(part)
    if skills:
        fields["must_have_skills"] = ", ".join(skills)

    for key in ALL_STEP_KEYS:
        fields.setdefault(key, "")

    return fields


def save_fields_to_session(fields: Dict[str, str]) -> None:
    """Persist fields in Streamlit session state."""
    st.session_state.setdefault("job_fields", {}).update(fields)


def display_fields_summary() -> None:
    """Show extracted fields as two-line bullet points."""

    fields = st.session_state.get("job_fields", {})
    lang = st.session_state.get("lang", "de")

    st.markdown(tr("### Extrahierte Jobdaten / Extracted Job Info", lang))
    st.markdown("<style>.field-bullet{font-size:16px;}</style>", unsafe_allow_html=True)
    for key, value in fields.items():
        if not value or key == "parsed_data_raw":
            continue
        st.markdown(
            f"- **{key.replace('_', ' ').title()}**<br>{value}",
            unsafe_allow_html=True,
        )
    if fields.get("parsed_data_raw"):
        st.text_area(
            label=tr("Parsed Data Raw", lang),
            value=fields["parsed_data_raw"],
            key="parsed_data_raw_summary",
            height=200,
        )


def display_fields_editable(prefix: str = "edit_") -> None:
    """Show all stored fields as editable inputs.

    Args:
        prefix: Key prefix to differentiate multiple widget groups.
    """

    fields = st.session_state.get("job_fields", {})
    lang = st.session_state.get("lang", "de")

    # Track used widget keys to avoid StreamlitDuplicateElementKey errors.
    used_keys = st.session_state.setdefault("_used_widget_keys", set())

    st.markdown(tr("### Extrahierte Jobdaten / Extracted Job Info", lang))
    for key, value in fields.items():
        if not value or key == "parsed_data_raw":
            continue
        widget_key = f"{prefix}{key}"
        if widget_key in used_keys:
            suffix = 1
            while f"{widget_key}_{suffix}" in used_keys:
                suffix += 1
            widget_key = f"{widget_key}_{suffix}"
        used_keys.add(widget_key)
        st.text_input(key.replace("_", " ").title(), value, key=widget_key)
    if fields.get("parsed_data_raw"):
        st.text_area(
            tr("Parsed Data Raw", lang),
            fields["parsed_data_raw"],
            key=f"{prefix}parsed_data_raw",
            height=200,
        )


def export_fields_as_markdown() -> None:
    """Provide a download link for the stored fields as Markdown."""
    fields = st.session_state.get("job_fields", {})
    lang = st.session_state.get("lang", "de")
    md = "\n".join(
        f"**{k.replace('_', ' ').capitalize()}:** {v}" for k, v in fields.items() if v
    )
    b64 = base64.b64encode(md.encode()).decode()
    href = (
        f'<a href="data:text/markdown;base64,{b64}" download="jobinfo.md">'
        f"{tr('Markdown herunterladen / Download Markdown', lang)}</a>"
    )
    st.markdown(href, unsafe_allow_html=True)


def display_all_fields_multiline_copy() -> None:
    """Show fields as multiline text areas."""
    fields = st.session_state.get("job_fields", {})
    lang = st.session_state.get("lang", "de")
    st.markdown(tr("### Alle Felder / All Fields", lang))
    for key, value in fields.items():
        if not value:
            continue
        st.text_area(key.replace("_", " ").title(), value, key=f"multi_{key}")
