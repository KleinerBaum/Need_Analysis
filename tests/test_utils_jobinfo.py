from utils.utils_jobinfo import basic_field_extraction, extract_text
from docx import Document


def test_basic_field_extraction_skills():
    text = "Proficiency in Python and machine learning libraries (e.g., scikit-learn, TensorFlow)."
    result = basic_field_extraction(text)
    assert (
        result["must_have_skills"]
        == "Python, machine learning libraries, scikit-learn, TensorFlow"
    )
    assert result["parsed_data_raw"] == text


def test_extract_text_from_docx_with_table(tmp_path):
    doc = Document()
    doc.add_paragraph("Jobtitel: Data Scientist")
    table = doc.add_table(rows=1, cols=1)
    table.cell(0, 0).text = (
        "Proficiency in Python and machine learning libraries (e.g., scikit-learn, TensorFlow)."
    )
    file_path = tmp_path / "sample.docx"
    doc.save(file_path)

    with file_path.open("rb") as fh:
        text = extract_text(fh)

    fields = basic_field_extraction(text)
    assert fields["job_title"] == "Data Scientist"
    assert (
        fields["must_have_skills"]
        == "Python, machine learning libraries, scikit-learn, TensorFlow"
    )
    assert fields["parsed_data_raw"] == text
